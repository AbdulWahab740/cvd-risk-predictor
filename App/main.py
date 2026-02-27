"""
Streamlit App - CVD Clinical Decision Support System
Loads the XGBoost model directly and queries Elasticsearch Agent Builder for recommendations.
"""

import streamlit as st
import requests
import json
import io
import os
import joblib
import numpy as np
import pandas as pd
import plotly.graph_objects as go
from datetime import datetime, timezone

try:
    import shap

    SHAP_AVAILABLE = True
except ImportError:
    SHAP_AVAILABLE = False

# ── Page config (must be first Streamlit call) ──────────────────────────────
st.set_page_config(
    page_title="CVD Risk Assessment",
    page_icon="🫀",
    layout="wide",
)

# ── Model loading ────────────────────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MODEL_PATH = os.path.join(BASE_DIR, "output", "cost_sensitive_xgboost_model.pkl")


@st.cache_resource(show_spinner="Loading ML model…")
def load_model():
    try:
        data = joblib.load(MODEL_PATH)
        mdl = data["model"]
        feats = data["features"]
        exp = shap.TreeExplainer(mdl) if SHAP_AVAILABLE else None
        return mdl, feats, exp
    except Exception as e:
        st.error(f"Could not load model: {e}")
        return None, [], None


model, features, explainer = load_model()

# ── Secrets ──────────────────────────────────────────────────────────────────
ELASTIC_API_KEY = st.secrets.get("ELASTIC_API_KEY", "")
ELASTIC_BASE_URL = st.secrets.get(
    "ELASTIC_BASE_URL",
    "https://my-elasticsearch-project-f43eb6.kb.asia-southeast1.gcp.elastic.cloud",
)
AGENT_ID = st.secrets.get("AGENT_ID", "cvd_7")

# ── Helper functions ─────────────────────────────────────────────────────────
FEATURE_NAMES = {
    "gender": "Gender",
    "age_years": "Age",
    "height": "Height",
    "weight": "Weight",
    "bmi": "Body Mass Index",
    "ap_hi": "Systolic Blood Pressure",
    "ap_lo": "Diastolic Blood Pressure",
    "pulse_pressure": "Pulse Pressure",
    "map": "Mean Arterial Pressure",
    "cholesterol": "Cholesterol Level",
    "gluc": "Glucose Level",
    "smoke": "Smoking Status",
    "alco": "Alcohol Consumption",
    "active": "Physical Activity",
}

FEATURE_UNITS = {
    "age_years": "years",
    "height": "cm",
    "weight": "kg",
    "bmi": "kg/m²",
    "ap_hi": "mmHg",
    "ap_lo": "mmHg",
    "pulse_pressure": "mmHg",
    "map": "mmHg",
    "cholesterol": "category",
    "gluc": "category",
    "smoke": "binary",
    "alco": "binary",
    "active": "binary",
    "gender": "binary",
}


def interpret(feature, value):
    if feature == "ap_hi":
        if value < 120:
            return "Normal systolic BP"
        elif value < 130:
            return "Elevated systolic BP"
        elif value < 140:
            return "Stage 1 hypertension"
        return "Stage 2 hypertension"
    if feature == "ap_lo":
        if value < 80:
            return "Normal diastolic BP"
        elif value < 90:
            return "Stage 1 hypertension"
        return "Stage 2 hypertension"
    if feature == "bmi":
        if value < 18.5:
            return "Underweight"
        elif value < 25:
            return "Normal weight"
        elif value < 30:
            return "Overweight"
        return "Obese"
    if feature == "pulse_pressure":
        return "Widened pulse pressure" if value > 60 else "Normal pulse pressure"
    if feature == "map":
        return "Elevated MAP" if value > 100 else "Normal MAP"
    if feature == "cholesterol":
        return {1: "Normal", 2: "Above normal", 3: "Well above normal"}.get(
            value, "Unknown"
        )
    if feature == "gluc":
        return {
            1: "Normal",
            2: "Above normal (prediabetes)",
            3: "Well above normal",
        }.get(value, "Unknown")
    if feature == "smoke":
        return "Active smoker" if value == 1 else "Non-smoker"
    if feature == "alco":
        return "Consumes alcohol" if value == 1 else "No alcohol"
    if feature == "active":
        return "Physically active" if value == 1 else "Sedentary"
    if feature == "gender":
        return "Female" if value == 1 else "Male"
    if feature == "age_years":
        if value < 40:
            return "Young adult"
        elif value < 60:
            return "Middle-aged"
        return "Older adult"
    return str(value)


def run_prediction(patient: dict, threshold: float = 0.5, top_n: int = 5):
    """Run local ML prediction and return structured result."""
    if model is None:
        raise RuntimeError("Model not loaded")

    input_dict = {f: patient.get(f) for f in features}
    df = pd.DataFrame([input_dict])[features]

    risk_score = float(model.predict_proba(df)[0, 1])
    prediction = 1 if risk_score >= threshold else 0
    risk_category = "HIGH" if prediction == 1 else "LOW"

    shap_vals = None
    if SHAP_AVAILABLE and explainer is not None:
        raw = explainer.shap_values(df)
        shap_vals = raw[1][0] if isinstance(raw, list) else raw[0]

    feat_importance = dict(zip(features, model.feature_importances_))
    all_factors = []
    for i, feat in enumerate(features):
        val = input_dict[feat]
        sv = float(shap_vals[i]) if shap_vals is not None else 0.0
        if val is None:
            val = 0
        all_factors.append(
            {
                "feature": feat,
                "feature_name": FEATURE_NAMES.get(feat, feat),
                "value": float(val) if isinstance(val, float) else int(val),
                "unit": FEATURE_UNITS.get(feat, ""),
                "importance": float(feat_importance[feat]),
                "shap_value": sv,
                "impact": "increases risk" if sv > 0 else "decreases risk",
                "magnitude": abs(sv),
                "interpretation": interpret(feat, val),
            }
        )

    top_factors = sorted(all_factors, key=lambda x: x["magnitude"], reverse=True)[
        :top_n
    ]
    for f in top_factors:
        del f["magnitude"]

    return {
        "risk_score": round(risk_score, 4),
        "risk_percentage": int(risk_score * 100),
        "prediction": int(prediction),
        "risk_category": risk_category,
        "top_risk_factors": top_factors,
        "model_confidence": round(max(risk_score, 1 - risk_score), 4),
        "prediction_timestamp": datetime.now(timezone.utc).isoformat(),
    }


def call_agent(prompt: str) -> str:
    """Call Elasticsearch Agent Builder and return the assistant message text."""
    if not ELASTIC_API_KEY:
        return "_No ELASTIC_API_KEY configured — agent recommendations unavailable._"

    headers = {
        "Authorization": f"ApiKey {ELASTIC_API_KEY}",
        "Content-Type": "application/json",
        "kbn-xsrf": "true",
    }
    payload = {"agent_id": AGENT_ID, "input": prompt}
    resp = requests.post(
        f"{ELASTIC_BASE_URL}/api/agent_builder/converse",
        headers=headers,
        json=payload,
        timeout=300,
    )
    resp.raise_for_status()
    data = resp.json()
    # Try common response shapes
    if "response" in data:
        r = data["response"]
        if isinstance(r, dict):
            return r.get("message", json.dumps(r, indent=2))
        return str(r)
    return json.dumps(data, indent=2)


# ── CSS ───────────────────────────────────────────────────────────────────────
st.markdown(
    """
<style>
.main-header{font-size:2.5rem;font-weight:bold;color:#1f77b4;margin-bottom:1rem}
.risk-high{background:#ffebee;color:#000;padding:1rem;border-left:5px solid #f44336;border-radius:5px}
.risk-moderate{background:#fff3e0;color:#000;padding:1rem;border-left:5px solid #ff9800;border-radius:5px}
.risk-low{background:#e8f5e9;color:#000;padding:1rem;border-left:5px solid #4caf50;border-radius:5px}
</style>
""",
    unsafe_allow_html=True,
)

# ── Header ────────────────────────────────────────────────────────────────────
st.markdown(
    '<div class="main-header">🫀 Cardiovascular Risk Assessment</div>',
    unsafe_allow_html=True,
)
st.markdown("AI-powered clinical decision support with evidence-based recommendations")

# ── Sidebar inputs ────────────────────────────────────────────────────────────
st.sidebar.header("📋 Patient Information")

with st.sidebar:
    st.subheader("Demographics")

    gender = st.selectbox(
        "Gender",
        options=[1, 2],
        format_func=lambda x: "Female" if x == 1 else "Male",
        help="Biological sex",
    )

    age_years = st.number_input(
        "Age (years)", min_value=18, max_value=100, value=55, step=1
    )

    col1, col2 = st.columns(2)
    with col1:
        height = st.number_input(
            "Height (cm)", min_value=140, max_value=220, value=170, step=1
        )
    with col2:
        weight = st.number_input(
            "Weight (kg)", min_value=40.0, max_value=200.0, value=80.0, step=0.5
        )

    height_m = height / 100
    bmi = weight / (height_m**2)
    st.metric("Calculated BMI", f"{bmi:.1f}")

    st.divider()
    st.subheader("Blood Pressure")

    col1, col2 = st.columns(2)
    with col1:
        ap_hi = st.number_input(
            "Systolic (mmHg)",
            min_value=80,
            max_value=250,
            value=140,
            step=5,
            help="Upper number",
        )
    with col2:
        ap_lo = st.number_input(
            "Diastolic (mmHg)",
            min_value=50,
            max_value=150,
            value=90,
            step=5,
            help="Lower number",
        )

    pulse_pressure = ap_hi - ap_lo
    map_value = (ap_hi + 2 * ap_lo) / 3
    st.caption(f"Pulse Pressure: {pulse_pressure} mmHg")
    st.caption(f"MAP: {map_value:.1f} mmHg")

    st.divider()
    st.subheader("Lab Values")

    cholesterol = st.radio(
        "Cholesterol",
        options=[1, 2, 3],
        format_func=lambda x: ["Normal", "Above Normal", "Well Above Normal"][x - 1],
        horizontal=True,
    )

    gluc = st.radio(
        "Glucose",
        options=[1, 2, 3],
        format_func=lambda x: ["Normal", "Above Normal", "Well Above Normal"][x - 1],
        horizontal=True,
    )

    st.divider()
    st.subheader("Lifestyle Factors")

    smoke = st.checkbox("Current Smoker", value=False)
    alco = st.checkbox("Consumes Alcohol", value=False)
    active = st.checkbox("Physically Active", value=True)

# ── Session state ─────────────────────────────────────────────────────────────
if "risk_result" not in st.session_state:
    st.session_state.risk_result = None
if "agent_text" not in st.session_state:
    st.session_state.agent_text = None
if "agent_prompt" not in st.session_state:
    st.session_state.agent_prompt = None

# ── Assess button ─────────────────────────────────────────────────────────────
if st.sidebar.button("🔍 Assess Risk", type="primary", use_container_width=True):
    patient_data = {
        "gender": gender,
        "age_years": float(age_years),
        "height": float(height),
        "weight": float(weight),
        "bmi": round(bmi, 1),
        "ap_hi": float(ap_hi),
        "ap_lo": float(ap_lo),
        "pulse_pressure": float(pulse_pressure),
        "map": round(map_value, 1),
        "cholesterol": cholesterol,
        "gluc": gluc,
        "smoke": 1 if smoke else 0,
        "alco": 1 if alco else 0,
        "active": 1 if active else 0,
    }

    with st.spinner("Calculating risk score…"):
        try:
            result = run_prediction(patient_data)
            st.session_state.risk_result = result
            st.session_state.agent_text = None  # clear stale recommendations
        except Exception as e:
            st.error(f"Prediction failed: {e}")
            st.stop()

    # Build and store prompt for the separate recommendations step
    st.session_state.agent_prompt = f"""Assess cardiovascular disease risk for:
- {age_years}-year-old {"female" if gender == 1 else "male"}
- Height: {height} cm, Weight: {weight} kg (BMI: {bmi:.1f})
- Blood Pressure: {ap_hi}/{ap_lo} mmHg
- Cholesterol: {["normal", "above normal", "well above normal"][cholesterol - 1]}
- Glucose: {["normal", "above normal", "well above normal"][gluc - 1]}
- Smoker: {"Yes" if smoke else "No"}
- Alcohol: {"Yes" if alco else "No"}
- Physically Active: {"Yes" if active else "No"}
- ML Risk Score: {result["risk_percentage"]}% ({result["risk_category"]} RISK)

Please provide a complete evidence-based clinical risk assessment with management recommendations."""
    st.rerun()

# ── Results display ───────────────────────────────────────────────────────────
if st.session_state.risk_result:
    result = st.session_state.risk_result

    risk_percentage = result["risk_percentage"]
    risk_category = result["risk_category"]
    top_factors = result["top_risk_factors"]

    # Risk gauge
    st.subheader("Risk Score")
    fig = go.Figure(
        go.Indicator(
            mode="gauge+number",
            value=risk_percentage,
            domain={"x": [0, 1], "y": [0, 1]},
            title={"text": "CVD Risk Score (%)", "font": {"size": 22}},
            gauge={
                "axis": {"range": [0, 100], "tickwidth": 1},
                "bar": {"color": "darkblue"},
                "bgcolor": "white",
                "borderwidth": 2,
                "bordercolor": "gray",
                "steps": [
                    {"range": [0, 40], "color": "#4caf50"},
                    {"range": [40, 70], "color": "#ff9800"},
                    {"range": [70, 100], "color": "#f44336"},
                ],
                "threshold": {
                    "line": {"color": "red", "width": 4},
                    "thickness": 0.75,
                    "value": 70,
                },
            },
        )
    )
    fig.update_layout(height=300)
    st.plotly_chart(fig, use_container_width=True)

    # Risk category banner
    if risk_category == "HIGH" or risk_percentage >= 70:
        st.markdown(
            f'<div class="risk-high"><b>⚠️ HIGH RISK</b> — {risk_percentage}% CVD Risk</div>',
            unsafe_allow_html=True,
        )
    elif risk_percentage >= 40:
        st.markdown(
            f'<div class="risk-moderate"><b>⚡ MODERATE RISK</b> — {risk_percentage}% CVD Risk</div>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            f'<div class="risk-low"><b>✅ LOW RISK</b> — {risk_percentage}% CVD Risk</div>',
            unsafe_allow_html=True,
        )

    st.divider()

    # Top risk factors
    st.subheader("Primary Risk Factors")
    for i, factor in enumerate(top_factors, 1):
        with st.expander(f"{i}. {factor['feature_name']} — {factor['interpretation']}"):
            c1, c2 = st.columns(2)
            with c1:
                st.metric("Value", f"{factor['value']} {factor['unit']}".strip())
                st.metric("Model Importance", f"{factor['importance'] * 100:.1f}%")
            with c2:
                if SHAP_AVAILABLE:
                    st.metric("SHAP Value", f"{factor['shap_value']:.3f}")
                st.metric("Impact", factor["impact"].title())

    st.divider()

    # Agent recommendations
    st.subheader("Evidence-Based Recommendations")

    if st.session_state.agent_text:
        st.markdown(st.session_state.agent_text)
    else:
        st.info(
            "Click below to fetch evidence-based clinical recommendations from the AI agent."
        )
        if st.button("📚 Get AI Recommendations", type="primary"):
            with st.spinner("Querying clinical guidelines (may take up to 2 minutes)…"):
                try:
                    if st.session_state.agent_prompt:
                        rec = call_agent(st.session_state.agent_prompt)
                    else:
                        rec = "_Agent prompt not available._"
                except Exception as e:
                    rec = f"_Agent recommendations unavailable: {e}_"
                st.session_state.agent_text = rec
                st.rerun()

    st.divider()

    # Download report
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_heading("CVD Clinical Risk Report", 0)
        doc.add_paragraph(
            f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
        )
        doc.add_heading("Risk Score", level=1)
        doc.add_paragraph(f"{risk_percentage}% — {risk_category} RISK")
        doc.add_heading("Primary Risk Factors", level=1)
        for i, f in enumerate(top_factors, 1):
            doc.add_paragraph(
                f"{i}. {f['feature_name']}: {f['value']} {f['unit']} ({f['interpretation']}) — {f['impact']}",
                style="List Number",
            )
        doc.add_heading("Evidence-Based Recommendations", level=1)
        doc.add_paragraph(st.session_state.agent_text or "(not yet loaded)")
        bio = io.BytesIO()
        doc.save(bio)
        st.download_button(
            label="📥 Download Clinical Report (.docx)",
            data=bio.getvalue(),
            file_name=f"cvd_risk_report_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except ImportError:
        st.info("Install `python-docx` to enable report download.")

else:
    # Landing / instructions
    st.info("👈 Fill in patient information in the sidebar and click **Assess Risk**")

    st.subheader("How It Works")
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown("### 1️⃣ ML Prediction")
        st.write(
            "XGBoost model analyses 14 patient features with SHAP interpretability — runs locally, no external call needed."
        )
    with c2:
        st.markdown("### 2️⃣ Evidence Retrieval")
        st.write(
            "Elasticsearch Agent Builder searches clinical guidelines and research literature."
        )
    with c3:
        st.markdown("### 3️⃣ Recommendations")
        st.write("Synthesises evidence-based clinical recommendations with citations.")

    st.divider()
    st.subheader("Sample Assessment")
    st.code(
        """
RISK SCORE: 78% (HIGH RISK)

PRIMARY RISK FACTORS:
1. Systolic BP 165 mmHg  → increases risk  (35% importance)
2. Cholesterol well above normal → increases risk (28%)
3. Active smoking          → increases risk  (15%)

RECOMMENDATIONS:
1. BLOOD PRESSURE MANAGEMENT
   Target: <130/80 mmHg
   First-line: ACE inhibitor or ARB
   Evidence: 2017 ACC/AHA HTN Guideline, Class I

2. LIPID MANAGEMENT ...
3. SMOKING CESSATION ...
""",
        language="text",
    )
